import os
import subprocess
import re
from docx import Document


def to_txt(input_folder, output_folder):
    os.makedirs(output_folder, exist_ok=True)

    for filename in os.listdir(input_folder):
        src = os.path.join(input_folder, filename)
        stem = os.path.splitext(filename)[0]
        txt_path = os.path.join(output_folder, stem + ".txt")

        # Extraction logic
        if filename.endswith(".doc"):
            subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'docx', src, '--outdir', output_folder],
                check=True
            )
            docx_path = os.path.join(output_folder, stem + ".docx")
            doc = Document(docx_path)
            text = "\n".join(p.text for p in doc.paragraphs)
            os.remove(docx_path)

        elif filename.endswith(".docx"):
            doc = Document(src)
            text = "\n".join(p.text for p in doc.paragraphs)

        # --- Flexible Regex Header Removal ---
        lines = text.splitlines()
        header_index = -1
        
        # Matches a line that contains only = and spaces, 
        # requiring at least one = to avoid matching empty lines.
        pattern = re.compile(r"^[ =]*=[ =]*$")

        for i, line in enumerate(lines[:10]):
            if pattern.match(line):
                header_index = i
                break
        
        if header_index != -1:
            lines = lines[header_index + 1:]
        
        # Rejoin and clean
        text = "\n".join(lines)
        text = text.replace('/', ',').replace(',,', ',')

        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(text)


if __name__ == "__main__":
    to_txt("private/original", "private/txt")
    print("Processing complete. Headers removed (supports spaces in separator).")